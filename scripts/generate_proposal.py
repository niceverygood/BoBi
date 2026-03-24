#!/usr/bin/env python3
"""이종인 영업이사 인센티브 제안서 생성 스크립트"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), '..', 'docs', '이종인_영업이사_제안서.docx')

doc = Document()

# ── 기본 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Apple SD Gothic Neo'
font.size = Pt(10.5)

# 한글 폰트 설정
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), 'Apple SD Gothic Neo')
rPr.append(rFonts)

# 제목 스타일
for level in ['Heading 1', 'Heading 2', 'Heading 3']:
    hs = doc.styles[level]
    hs.font.name = 'Apple SD Gothic Neo'
    hs.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    hrPr = hs.element.get_or_add_rPr()
    hrFonts = hrPr.makeelement(qn('w:rFonts'), {})
    hrFonts.set(qn('w:eastAsia'), 'Apple SD Gothic Neo')
    hrPr.append(hrFonts)


def add_table(doc, headers, rows, col_widths=None):
    """스타일된 테이블 생성"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)

    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9.5)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


# ═══════════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('영업이사 제안서')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('BoBi (보비) × 이종인 이사')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('AI 보험비서 BoBi 영업 파트너십 및 인센티브 구조')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('주식회사 바틀\n대표이사 한승수\n2026년 3월 24일')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 1. 제안 개요
# ═══════════════════════════════════════════════════
doc.add_heading('1. 제안 개요', level=1)

doc.add_paragraph(
    '본 제안서는 AI 보험비서 \'BoBi(보비)\' 서비스의 영업 파트너십에 대한 '
    '인센티브 구조 및 협업 조건을 안내드리기 위해 작성되었습니다.'
)

doc.add_heading('이종인 이사님을 영업이사로 모시고자 하는 이유', level=2)

reasons = [
    '보험설계 현장 경력 15년의 깊은 업계 이해도',
    'GA 본부장·지점장 네트워크 보유',
    '보험 설계사들의 Pain Point를 직접 체감하고 계시는 분',
    '현장 영업(구두전략)과 조직 관리에 모두 강점',
]
for r in reasons:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 2. 서비스 소개
# ═══════════════════════════════════════════════════
doc.add_heading('2. BoBi(보비) 서비스 소개', level=1)

doc.add_paragraph(
    'BoBi(보비)는 보험설계사를 위한 AI 보험비서입니다. '
    '고객의 건강보험심평원 진료이력 PDF를 업로드하면 AI가 자동으로 '
    '고지사항 분석, 보장 분석 리포트, 리모델링 제안서까지 생성합니다.'
)

doc.add_heading('핵심 기능', level=2)
features = [
    ('AI 고지사항 분석', '심평원 PDF 업로드 → 3개월/1년/5년 기간별 자동 분류'),
    ('보장 분석 리포트', '현재 보험 보장을 적정 기준과 비교 분석'),
    ('리모델링 제안서', 'AI 기반 유지/해지/변경 판정, 신규 가입 추천'),
    ('보험 자동 조회', 'CODEF 연동으로 고객 보험 계약 자동 조회'),
    ('PDF 다운로드', '분석 결과를 리포트 PDF로 즉시 출력'),
]
for title_text, desc_text in features:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc_text)

doc.add_paragraph()
doc.add_paragraph(
    '설계사 1건 분석에 약 30분이 절약되며, 월 30건 분석 시 약 15시간의 업무 시간을 줄여줍니다.',
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 3. 요금제
# ═══════════════════════════════════════════════════
doc.add_heading('3. 요금제 구성', level=1)

doc.add_heading('3-1. 개인 플랜', level=2)
add_table(doc,
    ['플랜', '월 가격', '분석 건수', '주요 기능'],
    [
        ['무료 체험', '0원', '5건/월', 'AI 고지사항 분석'],
        ['베이직 (추천)', '19,900원', '30건/월', '+ 보장 분석, PDF, CODEF'],
        ['프로', '39,900원', '무제한', '+ 리모델링 제안서, 우선 지원'],
    ])

doc.add_heading('3-2. 팀 / GA 플랜', level=2)
add_table(doc,
    ['플랜', '월 가격', '포함 인원', '추가 1인', '인당 단가'],
    [
        ['팀', '99,000원', '5명', '15,000원', '19,800원'],
        ['비즈니스', '590,000원', '30명', '17,000원', '19,667원'],
        ['엔터프라이즈', '1,490,000원', '100명', '12,900원', '14,900원'],
    ])

doc.add_heading('3-3. GA 규모별 예상 비용', level=2)
add_table(doc,
    ['GA 규모', '추천 플랜', '월 비용', '인당 단가', '할인율'],
    [
        ['50명', '비즈니스', '930,000원', '18,600원', '53%'],
        ['100명', '엔터프라이즈', '1,490,000원', '14,900원', '63%'],
        ['200명', '엔터프라이즈', '2,780,000원', '13,900원', '65%'],
    ])

p = doc.add_paragraph()
run = p.add_run('※ 개인 프로 플랜(39,900원/인) 대비 할인율')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 4. 인센티브 구조
# ═══════════════════════════════════════════════════
doc.add_heading('4. 인센티브 구조', level=1)

p = doc.add_paragraph()
run = p.add_run('※ 기본급 없음 · 100% 성과 연동형')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)

doc.add_paragraph(
    '보험 설계사 수수료 체계(FYC+RC)와 동일한 구조를 채택했습니다. '
    '신규 확보 시 1회 인센티브 + 유지 고객에 대한 매월 반복 인센티브로 구성됩니다.'
)

doc.add_heading('① 신규 유료전환 인센티브 (첫 결제 시 1회)', level=2)
doc.add_paragraph('이종인 이사님의 영업 활동을 통해 유료 전환된 고객 1건당 아래 금액을 지급합니다.')
add_table(doc,
    ['플랜', '월 가격', '인센티브', '비율'],
    [
        ['베이직', '19,900원', '2,000원/건', '약 10%'],
        ['프로', '39,900원', '4,000원/건', '약 10%'],
        ['팀', '99,000원', '10,000원/건', '약 10%'],
        ['비즈니스', '590,000원', '59,000원/건', '약 10%'],
        ['엔터프라이즈', '1,490,000원', '149,000원/건', '약 10%'],
    ])

doc.add_heading('② 리커링 인센티브 (매월 반복 지급)', level=2)
doc.add_paragraph(
    '유료 결제가 유지되는 한, 해당 고객 월 매출의 3%를 매월 지급합니다. '
    '고객이 늘수록 리커링 수입이 누적되어 안정적인 수익원이 됩니다.'
)

p = doc.add_paragraph()
run = p.add_run('리커링 누적 예시 (기본 시나리오, 전환율 12%)')
run.bold = True

add_table(doc,
    ['시점', '유료 고객 수', '월 리커링 수입'],
    [
        ['Q1 말 (3월)', '24명', '약 2만원/월'],
        ['Q2 말 (6월)', '240명', '약 24만원/월'],
        ['Q3 말 (9월)', '720명', '약 72만원/월'],
        ['Q4 말 (12월)', '1,200명', '약 120만원/월'],
    ])

p = doc.add_paragraph()
run = p.add_run('→ Q4 기준, 리커링만으로 월 120만원 + 전환 인센 합산 안정 수입 확보')
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

doc.add_heading('③ 분기 목표 달성 보너스', level=2)
doc.add_paragraph('아래 누적 가입자 마일스톤 달성 시 보너스를 별도 지급합니다.')
add_table(doc,
    ['분기', '목표 (누적 가입자)', '달성 보너스', '초과달성 (120%+)'],
    [
        ['Q1 (1~3월)', '200명', '50만원', '100만원'],
        ['Q2 (4~6월)', '2,000명', '200만원', '400만원'],
        ['Q3 (7~9월)', '6,000명', '300만원', '600만원'],
        ['Q4 (10~12월)', '10,000명', '500만원', '1,000만원'],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 5. 예상 수입 시뮬레이션
# ═══════════════════════════════════════════════════
doc.add_heading('5. 예상 수입 시뮬레이션', level=1)

doc.add_heading('5-1. 시나리오별 연간 보상', level=2)
add_table(doc,
    ['항목', '보수적 (전환율 8%)', '기본 (12%)', '낙관적 (18%)'],
    [
        ['기본급', '0원', '0원', '0원'],
        ['① 전환 인센티브', '약 210만원', '약 320만원', '약 480만원'],
        ['② 리커링 (누적)', '약 420만원', '약 660만원', '약 960만원'],
        ['③ 분기 보너스', '약 500만원', '약 1,050만원', '약 2,100만원'],
        ['연 총보상', '약 1,130만원', '약 2,030만원', '약 3,540만원'],
        ['월 환산', '약 94만원', '약 169만원', '약 295만원'],
    ])

doc.add_heading('5-2. 분기별 월 수입 추이 (기본 시나리오)', level=2)
add_table(doc,
    ['분기', '월 전환 인센', '월 리커링', '보너스(분기)', '월 환산 합계'],
    [
        ['Q1', '~3만원', '~1만원', '50만원', '~약 21만원'],
        ['Q2', '~18만원', '~12만원', '200만원', '~약 97만원'],
        ['Q3', '~27만원', '~48만원', '300만원', '~약 175만원'],
        ['Q4', '~27만원', '~96만원', '500만원', '~약 290만원'],
    ])

p = doc.add_paragraph()
run = p.add_run(
    '💡 핵심 포인트: Q1은 초기 시딩 기간이라 수입이 적지만, '
    'Q3부터 리커링 누적으로 월 150만원 이상의 안정적인 수입이 확보됩니다. '
    '이는 보험 설계사 유지 수수료(RC)와 동일한 구조입니다.'
)
run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 6. 회사 매출/영업이익과 인센티브 분배
# ═══════════════════════════════════════════════════
doc.add_heading('6. 회사 매출·영업이익 전망과 인센티브 분배', level=1)

doc.add_paragraph(
    '아래는 바틀(BoBi)의 2026년 분기별 매출·영업이익 전망과 '
    '이종인 이사님에 대한 인센티브 분배를 시계열로 정리한 것입니다.'
)

doc.add_heading('6-1. 분기별 매출·영업이익·인센티브 (기본 시나리오, 전환율 12%)', level=2)
add_table(doc,
    ['구분', 'Q1 (1~3월)', 'Q2 (4~6월)', 'Q3 (7~9월)', 'Q4 (10~12월)', '연간 합계'],
    [
        ['누적 가입자', '200명', '2,000명', '6,000명', '10,000명', '-'],
        ['유료 고객', '24명', '240명', '720명', '1,200명', '-'],
        ['매출', '256만원', '2,434만원', '7,303만원', '12,172만원', '22,165만원'],
        ['운영비용 (API+인프라)', '310만원', '460만원', '900만원', '1,350만원', '3,020만원'],
        ['이종인 이사 인센티브', '59만원', '310만원', '601만원', '943만원', '1,913만원'],
        ['영업이익', '-113만원', '1,664만원', '5,802만원', '9,879만원', '17,232만원'],
        ['영업이익률', '-', '68%', '79%', '81%', '78%'],
    ])

doc.add_heading('6-2. 인센티브 분배 비율 분석', level=2)
doc.add_paragraph(
    '이종인 이사님의 인센티브가 회사 매출 및 영업이익 대비 '
    '어느 정도의 비중을 차지하는지 투명하게 공개합니다.'
)

add_table(doc,
    ['구분', 'Q1', 'Q2', 'Q3', 'Q4', '연간'],
    [
        ['매출 대비 인센티브 비율', '23%', '13%', '8%', '8%', '9%'],
        ['영업이익 대비 인센티브 비율', '-', '19%', '10%', '10%', '11%'],
    ])

p = doc.add_paragraph()
run = p.add_run(
    '→ 연간 기준, 매출의 9%, 영업이익의 11%를 영업이사 인센티브로 분배'
)
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

doc.add_paragraph()

doc.add_heading('6-3. 왜 공정한 분배인가?', level=2)

fairness_points = [
    ('일반 SaaS 영업 수수료 기준',
     'B2B SaaS 업계의 영업사원 커미션은 통상 매출의 10~20% 수준입니다. '
     '본 구조(매출의 9%)는 업계 하단이지만, 리커링 누적과 보너스를 합산하면 적정 수준입니다.'),
    ('영업이익 대비 분배',
     '안정기(Q3~Q4) 기준으로 영업이익의 10%만 인센티브로 분배하고, '
     '90%는 회사에 귀속됩니다. 이종인 이사님이 만들어내는 매출 대비 매우 합리적인 비용입니다.'),
    ('고정비 0원의 가치',
     '기본급이 없으므로 매출=0이면 인센티브도 0입니다. '
     '회사 입장에서 리스크가 전혀 없고, 성과에 정확히 비례하여 공정합니다.'),
    ('이사님이 없으면 매출도 없다',
     '영업이사 없이 1만 명의 사용자를 확보하는 것은 사실상 불가능합니다. '
     'Q4 영업이익 약 9,879만원 중 943만원(10%)의 분배는 '
     '이사님의 기여도 대비 오히려 보수적인 수준입니다.'),
]

for title_text, desc_text in fairness_points:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title_text}\n')
    run.bold = True
    p.add_run(desc_text)

doc.add_paragraph()

# 핵심 메시지 강조
p = doc.add_paragraph()
run = p.add_run(
    '정리하면, 이종인 이사님의 인센티브는 회사 영업이익의 약 10~11%로, '
    '이사님이 만들어낸 매출·이익에 비례하여 합리적으로 분배되는 구조입니다.\n\n'
    '회사가 성장할수록 이사님의 절대 보상도 함께 커지며, '
    '이는 장기적인 파트너십을 위한 Win-Win 설계입니다.'
)
run.font.size = Pt(10.5)
run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 7. 영업 목표 및 전략
# ═══════════════════════════════════════════════════
doc.add_heading('7. 영업 목표 및 전략', level=1)

doc.add_heading('7-1. 분기별 목표', level=2)
add_table(doc,
    ['분기', '목표 (누적)', '신규 확보', '핵심 전략'],
    [
        ['Q1 (1~3월)', '200명', '200명', '초기 세팅, 지인 네트워크'],
        ['Q2 (4~6월)', '2,000명', '1,800명', '구두전략 (현장 영업)'],
        ['Q3 (7~9월)', '6,000명', '4,000명', 'GA 무료배포 + 구두전략 지속'],
        ['Q4 (10~12월)', '10,000명', '4,000명', '유료 전환 집중'],
    ])

doc.add_heading('7-2. 영업 방식', level=2)
methods = [
    ('구두전략 (현장 영업)', 'GA 지점 방문, 설명회 개최, 실시간 데모'),
    ('GA 단체 도입', '본부장·지점장 대상 비즈니스/엔터프라이즈 플랜 제안'),
    ('무료 체험 배포', '프로모션 코드를 활용한 무료 체험 기간 제공'),
    ('온라인 확산', '설계사 커뮤니티, 카카오톡 단체방 홍보'),
]
for title_text, desc_text in methods:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title_text}: ')
    run.bold = True
    p.add_run(desc_text)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 7. 왜 이 구조가 좋은가
# ═══════════════════════════════════════════════════
doc.add_heading('8. 본 인센티브 구조의 장점', level=1)

add_table(doc,
    ['관점', '이점'],
    [
        ['이종인 이사님', '보험 수수료(FYC+RC)와 동일한 리커링 구조 → 익숙하고 예측 가능'],
        ['이종인 이사님', '고객이 늘수록 리커링 누적 → 사실상 안정급여 역할'],
        ['이종인 이사님', '실적에 따라 보상 상한 없음 → 높은 성과 = 높은 보상'],
        ['바틀 (회사)', '초기 고정비 0원 → 캐시 플로우 안전'],
        ['바틀 (회사)', '성과만큼만 비용 발생 → 리스크 최소화'],
        ['상호 이익', '고객 해지 시 리커링도 감소 → CS 품질 자연 관리'],
    ])

doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 8. 계약 조건
# ═══════════════════════════════════════════════════
doc.add_heading('9. 계약 조건 (요약)', level=1)

conditions = [
    ('계약 형태', '업무위탁 계약 (프리랜서형)'),
    ('계약 기간', '1년 (자동 갱신, 상호 합의 하에 해지 가능)'),
    ('보상 방식', '월 정산 (익월 10일 지급)'),
    ('비경쟁', '계약 기간 중 동종 서비스 영업 제한'),
    ('기밀 유지', '고객 정보 및 사업 정보 비밀 유지'),
    ('프로모션 코드', '관리자 페이지에서 직접 생성·관리 가능'),
]

for label, value in conditions:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}: ')
    run.bold = True
    p.add_run(value)

doc.add_paragraph()
doc.add_paragraph()

# ═══════════════════════════════════════════════════
# 마무리
# ═══════════════════════════════════════════════════
doc.add_heading('10. 마무리', level=1)

doc.add_paragraph(
    '이종인 이사님의 15년 보험 현장 경험과 BoBi의 AI 기술이 만나면, '
    '보험설계사 시장에서 독보적인 서비스를 만들어낼 수 있다고 확신합니다.'
)
doc.add_paragraph(
    '본 제안에 대해 궁금한 점이나 조율이 필요한 부분이 있으시면 '
    '편하게 말씀해 주세요. 함께 성장하는 파트너가 되기를 기대합니다.'
)

doc.add_paragraph()
doc.add_paragraph()

closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = closing.add_run('주식회사 바틀\n대표이사 한승수\n연락처: 010-2309-7443')
run.font.size = Pt(10)

doc.add_paragraph()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('본 제안서는 사전 협의 자료이며, 최종 조건은 상호 합의를 통해 확정됩니다.')
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ── 저장 ──
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f'✅ 제안서 생성 완료: {OUTPUT_PATH}')
